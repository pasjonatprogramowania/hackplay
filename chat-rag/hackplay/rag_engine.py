import google.generativeai as genai
from typing import List, Dict
import faiss
import numpy as np
import pickle
import os
from config import Config
from PyPDF2 import PdfReader
import docx
import zipfile
import tempfile
import re


class GeminiRAGEngine:
    def __init__(self):
        self.config = Config()
        genai.configure(api_key=self.config.GEMINI_API_KEY)

        self.model = genai.GenerativeModel(self.config.GEMINI_MODEL)
        self.vision_model = genai.GenerativeModel('gemini-2.0-flash-lite')

        self.index = None
        self.documents = []
        self.metadata = []
        self._initialized = False

        self.index_path = os.path.join(self.config.WORKING_DIR, "faiss_index.bin")
        self.docs_path = os.path.join(self.config.WORKING_DIR, "documents.pkl")
        self.meta_path = os.path.join(self.config.WORKING_DIR, "metadata.pkl")

    def initialize(self):
        if self._initialized:
            return

        print("üöÄ Inicjalizacja Gemini RAG Engine z FAISS...")
        os.makedirs(self.config.WORKING_DIR, exist_ok=True)

        if os.path.exists(self.index_path):
            print("üìÇ ≈Åadowanie istniejƒÖcego indeksu...")
            self.index = faiss.read_index(self.index_path)
            with open(self.docs_path, 'rb') as f:
                self.documents = pickle.load(f)
            with open(self.meta_path, 'rb') as f:
                self.metadata = pickle.load(f)
        else:
            self.index = faiss.IndexFlatL2(self.config.EMBEDDING_DIM)
            self.documents = []
            self.metadata = []

        self._initialized = True
        print(f"‚úÖ FAISS gotowy! Dokument√≥w w bazie: {len(self.documents)}")

    def _save_index(self):
        faiss.write_index(self.index, self.index_path)
        with open(self.docs_path, 'wb') as f:
            pickle.dump(self.documents, f)
        with open(self.meta_path, 'wb') as f:
            pickle.dump(self.metadata, f)

    def _embed_single_text(self, text: str, task_type: str = "retrieval_document") -> List[float]:
        """Embed JEDEN tekst - z opcjƒÖ task_type"""
        result = genai.embed_content(
            model=f"models/{self.config.GEMINI_EMBEDDING_MODEL}",
            content=text,
            task_type=task_type
        )
        # Obs≈Çuga r√≥≈ºnych struktur odpowiedzi
        if 'embedding' in result:
            emb = result['embedding']
            if isinstance(emb, dict) and 'values' in emb:
                return emb['values']
            elif isinstance(emb, list):
                return emb
        raise ValueError(f"Unexpected result structure: {result}")

    def _embed_texts(self, texts: List[str]) -> np.ndarray:
        """Embed listƒô tekst√≥w - PO JEDNYM"""
        embeddings = []
        total = len(texts)
        print(f"üî¢ Embedding {total} chunk√≥w...")

        for idx, text in enumerate(texts):
            if idx % 10 == 0:
                print(f"   Progress: {idx}/{total}")

            try:
                emb = self._embed_single_text(text, task_type="retrieval_document")
                embeddings.append(emb)
            except Exception as e:
                print(f"‚ùå Error at {idx}: {e}")
                embeddings.append([0.0] * self.config.EMBEDDING_DIM)

        print(f"‚úÖ Embedding done: {len(embeddings)} wektor√≥w")
        return np.array(embeddings, dtype='float32')

    def _chunk_text(self, text: str) -> List[str]:
        """Lepszy chunking - zachowuje kontekst"""
        chunks = []

        # Najpierw dziel po artyku≈Çach
        article_pattern = r'(Art\.\s*\d+[a-z]?\.)'
        parts = re.split(article_pattern, text)

        current_chunk = ""
        article_header = ""

        for i, part in enumerate(parts):
            part = part.strip()
            if not part:
                continue

            # Je≈õli to nag≈Ç√≥wek Art.
            if re.match(article_pattern, part):
                article_header = part
                continue

            # To tre≈õƒá artyku≈Çu
            paragraphs = part.split('\n\n')

            for para in paragraphs:
                para = para.strip()
                if not para or len(para) < 20:
                    continue

                # Dodaj do chunka z nag≈Ç√≥wkiem
                test_chunk = f"{article_header}\n{current_chunk}\n{para}".strip()

                if len(test_chunk) < self.config.CHUNK_SIZE:
                    current_chunk = f"{current_chunk}\n{para}".strip()
                else:
                    # Zapisz chunk
                    if current_chunk:
                        final_chunk = f"{article_header}\n{current_chunk}".strip()
                        if len(final_chunk) > 100:
                            chunks.append(final_chunk)
                    current_chunk = para

            # Zapisz pozosta≈Çy chunk z artyku≈Çu
            if current_chunk:
                final_chunk = f"{article_header}\n{current_chunk}".strip()
                if len(final_chunk) > 100:
                    chunks.append(final_chunk)
                current_chunk = ""

        return chunks if chunks else [text[:self.config.CHUNK_SIZE]]

    def _extract_text_from_file(self, filepath: str) -> str:
        _, ext = os.path.splitext(filepath.lower())
        ext = ext.lstrip('.')  # remove leading dot

        if ext == 'pdf':
            reader = PdfReader(filepath)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text

        elif ext == 'docx':
            try:
                doc = docx.Document(filepath)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                        text += "\n"
                return text
            except Exception as e:
                return f"Error parsing DOCX: {str(e)}"

        elif ext == 'doc':
            try:
                doc = docx.Document(filepath)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except:
                return "Error: DOC not supported. Convert to DOCX."

        elif ext == 'txt':
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

        elif ext == 'zip':
            text = ""
            try:
                with zipfile.ZipFile(filepath, 'r') as zip_ref:
                    temp_dir = tempfile.mkdtemp()
                    zip_ref.extractall(temp_dir)
                    for root, dirs, files in os.walk(temp_dir):
                        for file in files:
                            file_path = os.path.join(root, file)
                            try:
                                text += f"\n\n--- {file} ---\n"
                                text += self._extract_text_from_file(file_path)
                            except:
                                continue
            except Exception as e:
                return f"Error extracting ZIP: {str(e)}"
            return text

        else:
            return f"Unsupported format: {ext}"

    def ingest_document(self, filepath: str, doc_id: str = None, metadata: Dict = None):
        if not self._initialized:
            self.initialize()

        print(f"üìÑ Przetwarzanie: {filepath}")
        text = self._extract_text_from_file(filepath)

        if text.startswith("Error") or text.startswith("Unsupported"):
            print(f"‚ùå {text}")
            return {"status": "error", "file": filepath, "error": text}

        chunks = self._chunk_text(text)
        print(f"üìù Chunk√≥w: {len(chunks)}")

        if len(chunks) == 0:
            return {"status": "error", "file": filepath, "error": "No text extracted"}

        print("üî¢ Generowanie embeddings...")
        embeddings = self._embed_texts(chunks)

        self.index.add(embeddings)

        base_meta = {
            "source": os.path.basename(filepath),
            "doc_id": doc_id or os.path.basename(filepath),
        }
        if metadata:
            base_meta.update(metadata)

        for i, chunk in enumerate(chunks):
            self.documents.append(chunk)
            chunk_meta = {**base_meta, "chunk_id": i}
            self.metadata.append(chunk_meta)

        self._save_index()

        print(f"‚úÖ Dodano {len(chunks)} chunk√≥w")
        return {"status": "success", "file": filepath, "chunks": len(chunks)}

    def query(self, question: str, top_k: int = None) -> Dict:
        if not self._initialized:
            self.initialize()

        top_k = top_k or self.config.TOP_K_RESULTS
        print(f"üîç Zapytanie: {question}")

        # Embed query - U≈ªYJ task_type="retrieval_query"!
        try:
            query_emb = self._embed_single_text(question, task_type="retrieval_query")
            query_embedding = np.array([query_emb], dtype='float32')
            print(f"‚úÖ Query embedding OK, shape: {query_embedding.shape}")
        except Exception as e:
            print(f"‚ùå Query embedding error: {e}")
            import traceback
            traceback.print_exc()
            return {"error": str(e), "answer": f"B≈ÇƒÖd: {e}", "sources": []}

        if self.index.ntotal == 0:
            return {
                "answer": "Brak dokument√≥w w bazie. Wrzuƒá pliki przez /api/upload",
                "sources": []
            }

        distances, indices = self.index.search(query_embedding, top_k)

        results = []
        context_chunks = []
        for idx, dist in zip(indices[0], distances[0]):
            if 0 <= idx < len(self.documents):
                results.append({
                    "chunk": self.documents[idx],
                    "metadata": self.metadata[idx],
                    "distance": float(dist)
                })
                context_chunks.append(self.documents[idx])

        if len(context_chunks) == 0:
            return {"answer": "Nie znaleziono pasujƒÖcych dokument√≥w.", "sources": []}

        context = "\n\n".join(context_chunks)

        prompt = f"""
        Jeste≈õ ekspertem prawnym specjalizujƒÖcym siƒô w polskim ustawodawstwie AI.
        Zachowuj sie jak czlowiek. Odpowiadaj zwiezle i na temat.
        KONTEKST z dokument√≥w ustawy o sztucznej inteligencji (UC71):
        {context}

        PYTANIE U≈ªYTKOWNIKA:
        {question}

        INSTRUKCJE:
        1. Odpowiedz TYLKO na podstawie dostarczonych dokument√≥w
        2. Cytuj konkretne artyku≈Çy (np. "Art. 67 ust. 1 stanowi...")
        3. Je≈õli informacji NIE MA w kontek≈õcie, napisz: "W dostarczonych fragmentach brak informacji o [X]"
        4. U≈ºywaj numeracji i punktor√≥w dla jasno≈õci
        5. Je≈õli pytanie dotyczy zmian, por√≥wnaj r√≥≈ºne wersje dokument√≥w

        ODPOWIED≈π:
        """

        try:
            response = self.model.generate_content(prompt)
            answer_text = response.text
        except Exception as e:
            print(f"‚ùå Generation error: {e}")
            answer_text = f"B≈ÇƒÖd: {e}"

        return {"answer": answer_text, "sources": results}

    def analyze_changes(self, doc_a_id: str, doc_b_id: str, context: str = None) -> Dict:
        if not self._initialized:
            self.initialize()

        chunks_a = [doc for doc, meta in zip(self.documents, self.metadata)
                    if meta.get('doc_id') == doc_a_id]
        chunks_b = [doc for doc, meta in zip(self.documents, self.metadata)
                    if meta.get('doc_id') == doc_b_id]

        text_a = "\n".join(chunks_a)
        text_b = "\n".join(chunks_b)

        prompt = f"""
Por√≥wnaj dokumenty:

DOK A ({doc_a_id}):
{text_a[:5000]}

DOK B ({doc_b_id}):
{text_b[:5000]}

{f"KONTEKST: {context}" if context else ""}

Zwr√≥ƒá:
1. KRYTYCZNE ZMIANY
2. ISTOTNE ZMIANY
3. NOWO≈öCI
4. REKOMENDACJE
"""

        response = self.vision_model.generate_content(prompt)

        return {"doc_a": doc_a_id, "doc_b": doc_b_id, "analysis": response.text}


rag_engine = GeminiRAGEngine()
